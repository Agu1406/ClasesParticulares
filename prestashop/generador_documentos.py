#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generador Automático de Documentos para Prácticas
==================================================
Este script genera documentos Word (.docx) a partir de archivos de código fuente
y descripciones de ejercicios/prácticas realizadas.

Uso:
    python generador_documentos.py

El script lee un archivo de configuración JSON donde defines:
- Información del trabajo (título, autor, fecha)
- Archivos de código a incluir
- Descripciones/explicaciones de cada ejercicio
- Formato y estructura del documento
"""

import json
import os
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def crear_documento_base():
    """Crea un documento Word con formato académico básico"""
    doc = Document()
    
    # Configurar estilo Normal (texto principal)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Lato'
    font.size = Pt(11)
    # Interlineado 1.15 (sencillo)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.15
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    
    # Configurar estilos de encabezados
    for i in range(1, 10):
        heading_style = doc.styles[f'Heading {i}']
        heading_font = heading_style.font
        heading_font.name = 'Lato'
        # Tamaños: H1=16, H2=14, H3=12, H4=11, resto=11
        sizes = {1: 16, 2: 14, 3: 12, 4: 11}
        heading_font.size = Pt(sizes.get(i, 11))
        heading_format = heading_style.paragraph_format
        heading_format.space_before = Pt(0)
        heading_format.space_after = Pt(0)
        heading_format.line_spacing = 1.15
    
    return doc

def agregar_titulo(doc, titulo, subtitulo=None):
    """Agrega un título al documento"""
    title = doc.add_heading(titulo, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Asegurar fuente Lato y tamaño
    for run in title.runs:
        run.font.name = 'Lato'
        run.font.size = Pt(18)
    title_format = title.paragraph_format
    title_format.space_before = Pt(0)
    title_format.space_after = Pt(0)
    title_format.line_spacing = 1.15
    
    if subtitulo:
        subtitle = doc.add_paragraph(subtitulo)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_format = subtitle.runs[0].font
        subtitle_format.name = 'Lato'
        subtitle_format.italic = True
        subtitle_format.size = Pt(12)
        subtitle.paragraph_format.space_before = Pt(0)
        subtitle.paragraph_format.space_after = Pt(0)
        subtitle.paragraph_format.line_spacing = 1.15
    
    # Solo un salto de línea, no párrafo vacío
    doc.add_paragraph()

def agregar_metadatos(doc, autor, fecha=None):
    """Agrega metadatos al documento"""
    if not fecha:
        fecha = datetime.now().strftime("%d/%m/%Y")
    
    p1 = doc.add_paragraph(f"Autor: {autor}")
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)
    p1.paragraph_format.line_spacing = 1.15
    for run in p1.runs:
        run.font.name = 'Lato'
        run.font.size = Pt(11)
    
    p2 = doc.add_paragraph(f"Fecha: {fecha}")
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    for run in p2.runs:
        run.font.name = 'Lato'
        run.font.size = Pt(11)
    
    doc.add_paragraph()  # Salto de línea

def agregar_codigo(doc, archivo_codigo, lenguaje="php"):
    """Agrega código fuente al documento con formato"""
    if not os.path.exists(archivo_codigo):
        doc.add_paragraph(f"[ERROR: No se encontró el archivo {archivo_codigo}]", style='Intense Quote')
        return
    
    # Título del archivo
    doc.add_heading(f"Archivo: {os.path.basename(archivo_codigo)}", level=2)
    
    # Leer contenido del archivo
    try:
        with open(archivo_codigo, 'r', encoding='utf-8') as f:
            codigo = f.read()
    except Exception as e:
        doc.add_paragraph(f"[ERROR al leer archivo: {str(e)}]", style='Intense Quote')
        return
    
    # Agregar código en un párrafo con formato monospace
    parrafo = doc.add_paragraph()
    run = parrafo.add_run(codigo)
    run.font.name = 'Consolas'  # Mantener Consolas para código
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)
    parrafo.paragraph_format.space_before = Pt(0)
    parrafo.paragraph_format.space_after = Pt(0)
    parrafo.paragraph_format.line_spacing = 1.15
    
    # Intentar aplicar formato de código (fondo gris)
    # Nota: Word no tiene un estilo de código nativo, usamos un estilo personalizado
    parrafo.style = 'Intense Quote'
    
    doc.add_paragraph()  # Salto de línea

def agregar_seccion(doc, titulo, contenido):
    """Agrega una sección con título y contenido"""
    heading = doc.add_heading(titulo, level=1)
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(0)
    heading.paragraph_format.line_spacing = 1.15
    
    if isinstance(contenido, list):
        for item in contenido:
            if isinstance(item, dict):
                # Si es un diccionario, puede tener subtítulo y texto
                if 'subtitulo' in item:
                    sub_heading = doc.add_heading(item['subtitulo'], level=2)
                    sub_heading.paragraph_format.space_before = Pt(0)
                    sub_heading.paragraph_format.space_after = Pt(0)
                    sub_heading.paragraph_format.line_spacing = 1.15
                if 'texto' in item:
                    p = doc.add_paragraph(item['texto'])
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.15
                    for run in p.runs:
                        run.font.name = 'Lato'
                        run.font.size = Pt(11)
                if 'codigo' in item:
                    agregar_codigo(doc, item['codigo'])
            else:
                p = doc.add_paragraph(str(item))
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    run.font.name = 'Lato'
                    run.font.size = Pt(11)
    else:
        p = doc.add_paragraph(str(contenido))
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        for run in p.runs:
            run.font.name = 'Lato'
            run.font.size = Pt(11)
    
    doc.add_paragraph()  # Salto de línea

def agregar_explicacion(doc, titulo, explicacion):
    """Agrega una explicación o descripción"""
    doc.add_heading(titulo, level=2)
    doc.add_paragraph(explicacion)
    doc.add_paragraph()  # Espacio

def agregar_imagen(doc, ruta_imagen, ancho=None, descripcion=None):
    """Agrega una imagen al documento"""
    if not os.path.exists(ruta_imagen):
        doc.add_paragraph(f"[ERROR: No se encontró la imagen {ruta_imagen}]", style='Intense Quote')
        return
    
    try:
        # Agregar imagen
        if ancho:
            doc.add_picture(ruta_imagen, width=Inches(ancho))
        else:
            # Ancho por defecto: 6 pulgadas (aprox. 15cm)
            doc.add_picture(ruta_imagen, width=Inches(6))
        
        # Agregar descripción si existe
        if descripcion:
            parrafo = doc.add_paragraph(descripcion)
            parrafo.style = 'Caption'
            parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            parrafo.paragraph_format.space_before = Pt(0)
            parrafo.paragraph_format.space_after = Pt(0)
            parrafo.paragraph_format.line_spacing = 1.15
            for run in parrafo.runs:
                run.font.name = 'Lato'
                run.font.size = Pt(10)  # Tamaño ligeramente menor para captions
        
        doc.add_paragraph()  # Salto de línea después de la imagen
    except Exception as e:
        doc.add_paragraph(f"[ERROR al insertar imagen: {str(e)}]", style='Intense Quote')

def agregar_indice(doc, puntos):
    """Agrega un índice automático al documento"""
    indice_heading = doc.add_heading("Índice", level=1)
    indice_heading.paragraph_format.space_before = Pt(0)
    indice_heading.paragraph_format.space_after = Pt(0)
    indice_heading.paragraph_format.line_spacing = 1.15
    
    for i, punto in enumerate(puntos, 1):
        p = doc.add_paragraph(f"{i}. {punto}", style='List Number')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        for run in p.runs:
            run.font.name = 'Lato'
            run.font.size = Pt(11)
    
    doc.add_page_break()  # Salto de página después del índice

def generar_documento(config_file='config_documento.json'):
    """Genera el documento Word a partir del archivo de configuración"""
    
    # Leer configuración
    if not os.path.exists(config_file):
        print(f"❌ Error: No se encontró el archivo de configuración: {config_file}")
        print("\n📝 Creando archivo de configuración de ejemplo...")
        crear_config_ejemplo(config_file)
        print(f"✅ Archivo creado: {config_file}")
        print("   Por favor, edita este archivo con tus datos y vuelve a ejecutar el script.")
        return
    
    with open(config_file, 'r', encoding='utf-8') as f:
        config = json.load(f)
    
    # Crear documento
    doc = crear_documento_base()
    
    # Agregar título
    titulo = config.get('titulo', 'Documento de Práctica')
    subtitulo = config.get('subtitulo', None)
    agregar_titulo(doc, titulo, subtitulo)
    
    # Agregar metadatos
    autor = config.get('autor', 'Estudiante')
    fecha = config.get('fecha', None)
    agregar_metadatos(doc, autor, fecha)
    
    # Agregar índice si hay ejercicios
    if 'ejercicios' in config:
        puntos_indice = [ej.get('titulo', f'Punto {i}') for i, ej in enumerate(config['ejercicios'], 1)]
        if 'conclusiones' in config:
            puntos_indice.append("Valoración Personal")
        agregar_indice(doc, puntos_indice)
    
    # Agregar introducción si existe
    if 'introduccion' in config:
        agregar_seccion(doc, "Introducción", config['introduccion'])
    
    # Agregar ejercicios/prácticas
    if 'ejercicios' in config:
        doc.add_heading("Ejercicios Realizados", level=1)
        
        for i, ejercicio in enumerate(config['ejercicios'], 1):
            # Título del ejercicio
            titulo_ej = ejercicio.get('titulo', f'Ejercicio {i}')
            ej_heading = doc.add_heading(titulo_ej, level=2)
            ej_heading.paragraph_format.space_before = Pt(0)
            ej_heading.paragraph_format.space_after = Pt(0)
            ej_heading.paragraph_format.line_spacing = 1.15
            
            # Descripción
            if 'descripcion' in ejercicio:
                p = doc.add_paragraph(ejercicio['descripcion'])
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    run.font.name = 'Lato'
                    run.font.size = Pt(11)
            
            # Archivos de código
            if 'archivos' in ejercicio:
                for archivo in ejercicio['archivos']:
                    ruta = archivo.get('ruta', '')
                    explicacion = archivo.get('explicacion', '')
                    
                    if explicacion:
                        p = doc.add_paragraph(explicacion, style='Intense Quote')
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.line_spacing = 1.15
                        for run in p.runs:
                            run.font.name = 'Lato'
                            run.font.size = Pt(11)
                    
                    agregar_codigo(doc, ruta)
            
            # Capturas de pantalla
            if 'capturas' in ejercicio:
                cap_heading = doc.add_heading("Capturas de Pantalla", level=3)
                cap_heading.paragraph_format.space_before = Pt(0)
                cap_heading.paragraph_format.space_after = Pt(0)
                cap_heading.paragraph_format.line_spacing = 1.15
                for captura in ejercicio['capturas']:
                    ruta_img = captura.get('ruta', '')
                    descripcion = captura.get('descripcion', '')
                    ancho = captura.get('ancho', 6)  # Ancho en pulgadas
                    
                    agregar_imagen(doc, ruta_img, ancho=ancho, descripcion=descripcion)
            
            # Explicación adicional
            if 'explicacion' in ejercicio:
                # Dividir por saltos de línea para mantener formato
                lineas = ejercicio['explicacion'].split('\n')
                for linea in lineas:
                    if linea.strip():  # Solo agregar si no está vacía
                        p = doc.add_paragraph(linea.strip())
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.line_spacing = 1.15
                        for run in p.runs:
                            run.font.name = 'Lato'
                            run.font.size = Pt(11)
                    else:
                        # Si es línea vacía, agregar párrafo vacío (salto de línea)
                        doc.add_paragraph()
            
            # Errores o problemas encontrados
            if 'errores' in ejercicio:
                error_heading = doc.add_heading("Problemas Encontrados", level=3)
                error_heading.paragraph_format.space_before = Pt(0)
                error_heading.paragraph_format.space_after = Pt(0)
                error_heading.paragraph_format.line_spacing = 1.15
                if isinstance(ejercicio['errores'], list):
                    for error in ejercicio['errores']:
                        p = doc.add_paragraph(f"• {error}", style='List Bullet')
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.line_spacing = 1.15
                        for run in p.runs:
                            run.font.name = 'Lato'
                            run.font.size = Pt(11)
                else:
                    p = doc.add_paragraph(ejercicio['errores'])
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.15
                    for run in p.runs:
                        run.font.name = 'Lato'
                        run.font.size = Pt(11)
            
            doc.add_paragraph()  # Espacio entre ejercicios
    
    # Agregar valoración personal si existe
    if 'valoracion_personal' in config:
        val_heading = doc.add_heading("Valoración Personal", level=1)
        val_heading.paragraph_format.space_before = Pt(0)
        val_heading.paragraph_format.space_after = Pt(0)
        val_heading.paragraph_format.line_spacing = 1.15
        if isinstance(config['valoracion_personal'], dict):
            # Si es un diccionario con preguntas específicas
            if 'como_ha_ido' in config['valoracion_personal']:
                h2 = doc.add_heading("¿Cómo os ha ido?", level=2)
                h2.paragraph_format.space_before = Pt(0)
                h2.paragraph_format.space_after = Pt(0)
                h2.paragraph_format.line_spacing = 1.15
                p = doc.add_paragraph(config['valoracion_personal']['como_ha_ido'])
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    run.font.name = 'Lato'
                    run.font.size = Pt(11)
            
            if 'mas_dificil' in config['valoracion_personal']:
                h2 = doc.add_heading("¿Qué habéis encontrado más difícil?", level=2)
                h2.paragraph_format.space_before = Pt(0)
                h2.paragraph_format.space_after = Pt(0)
                h2.paragraph_format.line_spacing = 1.15
                p = doc.add_paragraph(config['valoracion_personal']['mas_dificil'])
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    run.font.name = 'Lato'
                    run.font.size = Pt(11)
            
            if 'mas_interesante' in config['valoracion_personal']:
                h2 = doc.add_heading("¿Qué ha sido lo más interesante?", level=2)
                h2.paragraph_format.space_before = Pt(0)
                h2.paragraph_format.space_after = Pt(0)
                h2.paragraph_format.line_spacing = 1.15
                p = doc.add_paragraph(config['valoracion_personal']['mas_interesante'])
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    run.font.name = 'Lato'
                    run.font.size = Pt(11)
            
            if 'sugerencias' in config['valoracion_personal']:
                h2 = doc.add_heading("Sugerencias y Comentarios", level=2)
                h2.paragraph_format.space_before = Pt(0)
                h2.paragraph_format.space_after = Pt(0)
                h2.paragraph_format.line_spacing = 1.15
                p = doc.add_paragraph(config['valoracion_personal']['sugerencias'])
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    run.font.name = 'Lato'
                    run.font.size = Pt(11)
        else:
            # Si es texto directo
            p = doc.add_paragraph(config['valoracion_personal'])
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            for run in p.runs:
                run.font.name = 'Lato'
                run.font.size = Pt(11)
    
    # Agregar conclusiones si existen (para compatibilidad)
    if 'conclusiones' in config:
        agregar_seccion(doc, "Conclusiones", config['conclusiones'])
    
    # Guardar documento
    nombre_archivo = config.get('nombre_archivo', f'documento_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
    doc.save(nombre_archivo)
    
    print(f"✅ Documento generado exitosamente: {nombre_archivo}")
    return nombre_archivo

def crear_config_ejemplo(config_file='config_documento.json'):
    """Crea un archivo de configuración de ejemplo"""
    ejemplo = {
        "titulo": "Práctica 2 - Comercio Electrónico",
        "subtitulo": "Desarrollo de Tienda Online con PrestaShop",
        "autor": "Tu Nombre",
        "fecha": "06/12/2025",
        "nombre_archivo": "Practica2_ComercioElectronico.docx",
        "introduccion": [
            "Este documento describe el desarrollo de la Práctica 2 sobre Comercio Electrónico.",
            "Se ha implementado una tienda online utilizando PrestaShop con las siguientes funcionalidades:",
            "- Configuración inicial del sistema",
            "- Gestión de productos",
            "- Configuración de métodos de pago",
            "- Personalización del tema"
        ],
        "ejercicios": [
            {
                "titulo": "Ejercicio 1: Instalación de PrestaShop",
                "descripcion": "Se procedió a instalar PrestaShop en el servidor local siguiendo los pasos del enunciado.",
                "archivos": [
                    {
                        "ruta": "ruta/al/archivo/config.php",
                        "explicacion": "Archivo de configuración principal de PrestaShop"
                    }
                ],
                "explicacion": "La instalación se realizó correctamente y se configuraron los parámetros básicos del sistema."
            },
            {
                "titulo": "Ejercicio 2: Configuración de Productos",
                "descripcion": "Se configuraron los productos de la tienda con sus respectivas categorías.",
                "archivos": [
                    {
                        "ruta": "ruta/al/archivo/productos.php",
                        "explicacion": "Script para gestión de productos"
                    }
                ]
            }
        ],
        "conclusiones": [
            "La práctica se ha completado exitosamente.",
            "Se han implementado todas las funcionalidades requeridas.",
            "El sistema está funcionando correctamente y listo para producción."
        ]
    }
    
    with open(config_file, 'w', encoding='utf-8') as f:
        json.dump(ejemplo, f, ensure_ascii=False, indent=4)

if __name__ == "__main__":
    print("=" * 60)
    print("📄 Generador Automático de Documentos para Prácticas")
    print("=" * 60)
    print()
    
    try:
        generar_documento()
    except ImportError:
        print("❌ Error: No se encontró la librería 'python-docx'")
        print()
        print("📦 Para instalar las dependencias necesarias, ejecuta:")
        print("   pip install python-docx")
        print()
    except Exception as e:
        print(f"❌ Error inesperado: {str(e)}")
        import traceback
        traceback.print_exc()

